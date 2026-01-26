import pdfplumber
from docx import Document
from docx2pdf import convert
from copy import deepcopy
from docx.table import _Row,Table
from docx.oxml.text.run import CT_Text
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import platform
import time
import re
import unicodedata
import shutil
import os
import time
import threading
import pythoncom

# 1. UTILIDADES
def normalize_text(text):
    # Mantiene acentos y caracteres especiales
    text = unicodedata.normalize("NFC", text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{2,}', '\n', text)
    return text.strip()


def read_pdf(path):
    blocks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text(layout=True)
            if txt:
                blocks.append(txt)
    return "\n".join(blocks)


def reemplazar_variables_xml(xml_element, data):
    """
    Reemplaza placeholders SOLO en nodos CT_Text
    (únicos que permiten escritura segura)
    """
    for node in xml_element.iter():
        if isinstance(node, CT_Text) and node.text:
            for k, v in data.items():
                placeholder = f"{{{{{k}}}}}"
                if placeholder in node.text:
                    node.text = node.text.replace(placeholder, str(v))    


def iterar_bloques(doc):
    for child in doc.element.body:
        if child.tag.endswith('}p'):
            yield Paragraph(child, doc)
        elif child.tag.endswith('}tbl'):
            yield Table(child, doc)


def guardar_bloque_mixto(doc, inicio, fin):
    bloque = []
    capturar = False
    inicio_p = None
    fin_p = None
    buffer = []

    for elemento in iterar_bloques(doc):

        if isinstance(elemento, Paragraph):
            texto = elemento.text.strip()

            # --- INICIO ---
            if inicio in texto:
                capturar = True
                inicio_p = elemento
                continue

            # --- FIN ---
            if fin in texto and capturar:
                fin_p = elemento
                break

            if capturar:
                buffer.append(elemento)
                bloque.append(("paragraph", deepcopy(elemento._p)))

        elif isinstance(elemento, Table) and capturar:
            buffer.append(elemento)
            bloque.append(("table", deepcopy(elemento._tbl)))

    # -------------------------------
    # BORRADO REAL DEL DOCUMENTO
    # -------------------------------
    def delete_element(el):
        if isinstance(el, Paragraph) and el._p is not None:
            el._element.getparent().remove(el._element)
            el._p = el._element = None
        elif isinstance(el, Table):
            el._tbl.getparent().remove(el._tbl)

    # borrar contenido interno
    for el in buffer:
        delete_element(el)

    # borrar marcadores INICIO / FIN
    if inicio_p:
        delete_element(inicio_p)
    if fin_p:
        delete_element(fin_p)

    return bloque


def pegar_bloque_con_datos(doc, bloque, lista_datos):
    body = doc.element.body

    for datos in lista_datos:
        for _, xml in bloque:
            xml_clonado = deepcopy(xml)
            reemplazar_variables_xml(xml_clonado, datos)
            body.append(xml_clonado)


# 2. NORMALIZACION ATS
def rebuild_structure(text):
    headers = [
        "perfil profesional", "profile",
        "experiencia laboral", "work experience",
        "education", "educacion",
        "skills", "habilidades",
        "languages", "idiomas"
    ]

    for h in headers:
        text = re.sub(rf"\s*{h}\s*", f"\n{h.upper()}\n", text, flags=re.IGNORECASE)

    # Insertar salto de linea antes de cualquier • (u otros bullets) y mantener el simbolo
    text = re.sub(r"\s*([•\*\|▪●])\s*", r"\n\1 ", text)
    
    return normalize_text(text)

def normalize_experience_lines(lines):
    cleaned = []
    for l in lines:
        l = l.replace("|", "").strip()
        l = re.sub(r'^[●•\-]\s*', '', l)
        cleaned.append(l)
    return cleaned

# 3. FORMATEOS
DATE_REGEX = re.compile(
    r"""
    (
        # 03/2025 - 09/2025
        \d{2}/\d{4}\s*[-–]\s*(\d{2}/\d{4}|actualidad|present|current)
        |
        # 2013 - 2014
        \d{4}\s*[-–]\s*(\d{4}|actualidad|present|current)
        |
        # Mar 2015 - Sep 2017
        (ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic|
         enero|febrero|marzo|abril|mayo|junio|julio|agosto|
         septiembre|setiembre|octubre|noviembre|diciembre|
         jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)
        \s+\d{4}\s*[-–]\s*
        (
            (ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic|
             enero|febrero|marzo|abril|mayo|junio|julio|agosto|
             septiembre|setiembre|octubre|noviembre|diciembre|
             jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)
            \s+\d{4}
            |
            actualidad|present|current
        )
    )
    """,
    re.IGNORECASE | re.VERBOSE
)

def format_experiencia_bloques(bloques):
    salida = []
    for b in bloques:
        salida.append(
            f"""{b['fecha']}
Empresa: {b['empresa']}
Puesto: {b['puesto']}
Funciones:
""" + "\n".join(f"• {f}" for f in b["funciones"])
        )
    return "\n\n".join(salida)

def format_experiencia_plantilla(lines):
    bloques = []
    actual = None

    for l in lines:
        l = l.strip()
        if not l:
            continue

        # Fecha sola
        if DATE_REGEX.search(l) and len(l.split()) <= 6:
            if actual:
                actual["fecha"] = l
            continue

        # Puesto + fecha (EUROPASS)
        if DATE_REGEX.search(l) and " – " in l:
            if actual:
                actual["puesto"] = l.split(" – ")[0].strip()
                actual["fecha"] = " – ".join(l.split(" – ")[1:])
            continue

        # Empresa (EUROPASS ICONO / TEXTO)
        if "–" in l and any(city in l.lower() for city in ["madrid", "gijon", "oviedo", "spain"]):
            if actual:
                bloques.append(actual)
            actual = {
                "empresa": l.replace("", "").strip(),
                "puesto": "",
                "fecha": "",
                "funciones": []
            }
            continue

        # Eempresa clasica
        if l.isupper() and len(l.split()) <= 6:
            if actual:
                bloques.append(actual)
            actual = {
                "empresa": l,
                "puesto": "",
                "fecha": "",
                "funciones": []
            }
            continue

        if not actual:
            continue

        # Puesto
        if not actual["puesto"]:
            actual["puesto"] = l
            continue

        # Funciones
        actual["funciones"].append(l)

    if actual:
        bloques.append(actual)

    # Formato final
    salida = []
    for b in bloques:
        salida.append(
            f"""{b['fecha']}
Empresa: {b['empresa']}
Puesto: {b['puesto']}
Funciones:
""" + "\n".join(f"• {f}" for f in b["funciones"])
        )

    return "\n\n".join(salida)

def format_proyectos(lines):
    bloques = []
    actual = []

    for l in lines:
        # Titulo del proyecto
        if not l.startswith(("•", "-", "*")) and len(l.split()) <= 6:
            if actual:
                bloques.append("\n".join(actual))
                actual = []
            actual.append(l)
        else:
            actual.append(l)

    if actual:
        bloques.append("\n".join(actual))

    return "\n\n".join(bloques)

def clean_bullets(lines):
    return [re.sub(r'^[•\-\*]\s*', '', l) for l in lines]

def cv_json_to_docx_data(cv):
    return {
        "NOMBRE": cv.get("nombre", ""),
        "INICIALES": cv.get("nombreinis", ""),
        "EMAIL": cv.get("contacto", {}).get("email", ""),
        "TELEFONO": cv.get("contacto", {}).get("telefono", ""),
        "GITHUB": cv.get("contacto", {}).get("github", ""),
        "LINKEDIN": cv.get("contacto", {}).get("linkedin", ""),
        "UBI": cv.get("contacto", {}).get("provincia_pais", ""),

        "PERFIL": cv.get("perfil", ""),
        "RESUMEN": cv.get("resumen", ""),
        "RESUMENP": cv.get("resumen_profesional", ""),
        "ESPECIALIZACION": cv.get("areas_especializacion", ""),
        "ANYOS": cv.get("anyos", ""),

        "SKILLS": " | ".join(cv.get("skills", [])),
        #"FORMACION": "\n".join(cv.get("educacion", [])),
        "EDUCACION": cv.get("educacion", ""),
        "CERTIFICACIONES": "\n".join(cv.get("certificaciones", [])),

        "EXPERIENCIA_PLANTILLA": cv.get("experiencia", ""),

        "IDIOMAS": "\n".join(
            f"• {k}: {v}" for k, v in cv.get("idiomas", {}).items()
        ),

        "PROYECTOS": cv.get("proyectos_formateados", "")
    }

def is_empty_value(v):
    if v is None:
        return True
    if isinstance(v, str) and not v.strip():
        return True
    if isinstance(v, (list, dict)) and len(v) == 0:
        return True
    return 

def remove_inline_empty_fields(doc, data):
    def delete_paragraph(p):
        if p._p is not None:
            el = p._element
            el.getparent().remove(el)
            p._p = p._element = None

    for p in list(doc.paragraphs):
        if p._p is None:
            continue

        text = p.text.strip()

        for k, v in data.items():
            placeholder = f"{{{{{k}}}}}"

            # placeholder inline + valor vacío → borrar línea
            if placeholder in text and is_empty_value(v):
                delete_paragraph(p)
                break

def remove_empty_blocks(doc, data):
    def delete_paragraph(p):
        if p._p is not None:
            el = p._element
            el.getparent().remove(el)
            p._p = p._element = None

    paragraphs = list(doc.paragraphs)
    active_key = None
    block_paragraphs = []
    start_marker = None

    for p in paragraphs:
        if p._p is None:
            continue

        text = p.text.strip()

        # -------- INICIO BLOQUE --------
        if text.startswith("{{#") and text.endswith("}}"):
            active_key = text[3:-2]
            block_paragraphs = []
            start_marker = p
            continue

        # -------- FIN BLOQUE --------
        if text.startswith("{{/") and text.endswith("}}"):
            end_key = text[3:-2]

            # eliminar SIEMPRE marcadores
            delete_paragraph(start_marker)
            delete_paragraph(p)

            # si el bloque estaba vacío → borrar contenido interno
            if active_key == end_key and is_empty_value(data.get(active_key)):
                for bp in block_paragraphs:
                    delete_paragraph(bp)

            # reset
            active_key = None
            block_paragraphs = []
            start_marker = None
            continue

        # -------- CONTENIDO BLOQUE --------
        if active_key:
            block_paragraphs.append(p)

def replace_placeholders(doc, data, empty_text=""):
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    paragraphs = list(doc.paragraphs)

    for i, p in enumerate(paragraphs):
        full_text = p.text.strip()

        for k, v in data.items():
            placeholder = f"{{{{{k}}}}}"

            if placeholder not in full_text:
                continue

            # -------- CASO: VALOR VACÍO --------
            if is_empty_value(v):

                # 1️⃣ Si el párrafo SOLO contiene el placeholder
                if full_text == placeholder:
                    delete_paragraph(p)

                    # 2️⃣ Buscar hacia arriba el primer párrafo NO vacío
                    j = i - 1
                    while j >= 0:
                        prev = paragraphs[j]
                        if prev.text and prev.text.strip():
                            delete_paragraph(prev)
                            break
                        j -= 1

                    break

                # 3️⃣ Placeholder mezclado en texto
                else:
                    p.text = p.text.replace(placeholder, empty_text)

            # -------- CASO: CON VALOR --------
            else:
                p.text = p.text.replace(placeholder, str(v))

    # -------- TABLAS (sin borrar títulos) --------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    full_text = p.text.strip()

                    for k, v in data.items():
                        placeholder = f"{{{{{k}}}}}"

                        if placeholder not in full_text:
                            continue

                        if is_empty_value(v):
                            if full_text == placeholder:
                                delete_paragraph(p)
                                break
                            else:
                                p.text = p.text.replace(placeholder, empty_text)
                        else:
                            p.text = p.text.replace(placeholder, str(v))


def replace_placeholders_preserve_style(doc, data, empty_text=""):
    def replace_in_runs(runs, data):
        for run in runs:
            for k, v in data.items():
                placeholder = f"{{{{{k}}}}}"

                if placeholder not in run.text:
                    continue

                if is_empty_value(v):
                    run.text = run.text.replace(placeholder, empty_text)
                else:
                    run.text = run.text.replace(placeholder, str(v))

    # Parrafos
    for p in doc.paragraphs:
        replace_in_runs(p.runs, data)

    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_runs(p.runs, data)

def extract_format_blocks(doc):
    blocks = {}
    paragraphs = list(doc.paragraphs)

    active_key = None
    text_buffer = []
    para_buffer = []
    start_p = None

    def delete_paragraph(p):
        if p and p._p is not None:
            el = p._element
            el.getparent().remove(el)
            p._p = p._element = None

    # -------- PÁRRAFOS --------
    for p in paragraphs:
        if p._p is None:
            continue

        text = p.text.strip()

        if text.startswith("{{_") and text.endswith("}}"):
            active_key = text[3:-2]
            text_buffer = []
            para_buffer = []
            start_p = p
            continue

        if text.startswith("{{-") and text.endswith("}}") and active_key:
            end_key = text[3:-2]

            if end_key == active_key:
                blocks[active_key] = "\n".join(text_buffer)

                delete_paragraph(start_p)
                delete_paragraph(p)
                for bp in para_buffer:
                    delete_paragraph(bp)

            active_key = None
            text_buffer = []
            para_buffer = []
            start_p = None
            continue

        if active_key:
            text_buffer.append(p.text)
            para_buffer.append(p)

    # -------- TABLAS (NUEVO) --------
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells)

            if row_text.startswith("{{_") and row_text.endswith("}}"):
                key = row_text[3:-2]
                blocks[key] = row

    return blocks

def encontrar_parrafo_por_texto(doc, texto):
    for p in doc.paragraphs:
        if texto in p.text:
            return p
    return None

def generate_cv_from_template(
    template_path,
    cv_json,
    output_dir="output",
    plantilla_nombre="Plantilla"
    ):
    """
    Genera un DOCX y un PDF desde la plantilla usando docx2pdf.
    Usa threading + pythoncom.CoInitialize() para que funcione en Flask.
    """
    os.makedirs(output_dir, exist_ok=True)

    # Preparar nombres unicos
    def sanitize(text):
        return re.sub(r'[^A-Za-z0-9_-]', '', text.replace(" ", "_"))

    safe_name = sanitize(cv_json["nombre"] or "CV")
    safe_plantilla = sanitize(plantilla_nombre)
    timestamp = int(time.time())

    docx_out = os.path.join(
        output_dir,
        f"CV_{safe_name}_{safe_plantilla}_{timestamp}.docx"
    )

    pdf_out = os.path.join(
        output_dir,
        f"CV_{safe_name}_{safe_plantilla}_{timestamp}.pdf"
    )

    # Limpiar archivos antiguos
    if os.path.exists(docx_out):
        os.remove(docx_out)
    if os.path.exists(pdf_out):
        os.remove(pdf_out)

    # Copiar plantilla
    shutil.copy(template_path, docx_out)
    doc = Document(docx_out)

    format_blocks = extract_format_blocks(doc)
    cv_json["_formatos"] = format_blocks

    data = cv_json_to_docx_data(cv_json)

    # -------------------------------------------------
    # EXPERIENCIA (BLOQUE VISUAL, PEGADO EN PLACEHOLDER)
    # -------------------------------------------------
    experiencias = cv_json.get("experiencia", [])

    if experiencias and isinstance(experiencias[0], dict):

        bloque_exp = guardar_bloque_mixto(
            doc,
            "{{INICIO_EXPERIENCIA}}",
            "{{FINAL_EXPERIENCIA}}"
        )

        bloque_func = guardar_bloque_mixto(
            doc,
            "{{INICIO_FUNCION}}",
            "{{FINAL_FUNCION}}"
        )

        p_target = None
        for p in doc.paragraphs:
            if "{{EXPERIENCIA_PLANTILLA}}" in p.text:
                p_target = p
                break

        if bloque_exp and p_target:
            parent = p_target._element.getparent()
            idx = parent.index(p_target._element)
            parent.remove(p_target._element)

            for exp in experiencias:

    # -------------------------
    # CABECERA (UNA SOLA VEZ)
    # -------------------------
                for _, xml in bloque_exp:
                    xml_clonado = deepcopy(xml)

                    reemplazar_variables_xml(
                        xml_clonado,
                        {
                            "fecha_inicio": exp.get("fecha_inicio", ""),
                            "fecha_fin": exp.get("fecha_fin", ""),
                            "empresa": exp.get("empresa", ""),
                            "puesto": exp.get("puesto", ""),
                            "ubicacion": exp.get("ubicacion", ""),
                        }
                    )

                    parent.insert(idx, xml_clonado)
                    idx += 1

                # -------------------------
                # FUNCIONES (UNA DEBAJO DE OTRA)
                # -------------------------
                for funcion in exp.get("funciones", []):
                    for _, xml in bloque_func:
                        xml_clonado = deepcopy(xml)

                        reemplazar_variables_xml(
                            xml_clonado,
                            {"funcion": funcion}
                        )

                        parent.insert(idx, xml_clonado)
                        idx += 1

                    # -------------------------
                # FUNCIONES (UNA DEBAJO DE OTRA)
                # -------------------------
                for funcion in exp.get("funciones", []):
                    for _, xml in bloque_func:
                        xml_clonado = deepcopy(xml)

                        reemplazar_variables_xml(
                            xml_clonado,
                            {"funcion": funcion}
                        )

                        parent.insert(idx, xml_clonado)
                        idx += 1

                p_empty = OxmlElement("w:p")
                parent.insert(idx, p_empty)
                idx += 1

    # -------------------------------------------------
    # EDUCACION (BLOQUE VISUAL, PEGADO EN PLACEHOLDER)
    # -------------------------------------------------
    educaciones = cv_json.get("educacion", [])

    if educaciones and isinstance(educaciones[0], dict):

        bloque_edu = guardar_bloque_mixto(
            doc,
            "{{INICIO_EDUCACION}}",
            "{{FINAL_EDUCACION}}"
        )

        p_target = None
        for p in doc.paragraphs:
            if "{{EDUCACION}}" in p.text:
                p_target = p
                break

        if bloque_edu and p_target:
            parent = p_target._element.getparent()
            idx = parent.index(p_target._element)
            parent.remove(p_target._element)

            for ed in educaciones:
                for _, xml in bloque_edu:
                    xml_clonado = deepcopy(xml)
                    reemplazar_variables_xml(
                        xml_clonado,
                            {
                                "fecha_inicio": ed.get("fecha_inicio", ""),
                                "fecha_fin": ed.get("fecha_fin", ""),
                                "titulo": ed.get("titulo", ""),
                                "institucion": ed.get("institucion", ""),
                                "ubicacion": ed.get("ubicacion", ""),
                                "nota_final": ed.get("nota_final", "")
                            }
                    )
                    parent.insert(idx, xml_clonado)
                    idx += 1

    # -------------------------------------------------
    # REEMPLAZO GENERAL + LIMPIEZA LEGACY
    # -------------------------------------------------
    replace_placeholders_preserve_style(doc, data)
    time.sleep(1)

    # SOLO usar remove_empty_blocks si NO hay bloques visuales
    remove_empty_blocks(doc, data)
    time.sleep(1)

    remove_inline_empty_fields(doc, data)

    doc.save(docx_out)

    # -------------------------------------------------
    # PDF
    # -------------------------------------------------
    pdf_generated = False

    def convert_pdf_thread():
        nonlocal pdf_generated
        try:
            pythoncom.CoInitialize()
            convert(docx_out, pdf_out)
            if os.path.exists(pdf_out):
                pdf_generated = True
        except Exception as e:
            print("Error generando PDF en hilo:", e)

    if platform.system().lower() == "windows":
        thread = threading.Thread(target=convert_pdf_thread)
        thread.start()
        thread.join()

        if not pdf_generated:
            pdf_out = None
    else:
        pdf_out = None

    return docx_out, pdf_out if pdf_generated else None
